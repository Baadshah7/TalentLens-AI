import os
import re
import json
import pdfplumber
import docx
import spacy
from typing import Dict, Any, List, Tuple, Optional
from ethical import sanitize_text

# Load spaCy model
try:
    nlp = spacy.load("en_core_web_sm")
except Exception:
    # Fallback if model load fails in some environment
    nlp = None

# Regex patterns
EMAIL_REGEX = re.compile(r'[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+')
PHONE_REGEX = re.compile(r'\+?\d{1,4}?[-.\s]?\(?\d{1,3}?\)?[-.\s]?\d{1,4}[-.\s]?\d{1,4}[-.\s]?\d{1,9}')
YEAR_REGEX = re.compile(r'\b(19\d{2}|20\d{2})\b')
DATE_RANGE_REGEX = re.compile(
    r'(\b(?:0[1-9]|1[0-2]|[A-Za-z]{3,9})[\s/.-]*(?:20\d{2}|19\d{2})|\b(?:19|20)\d{2})\s*[–—\-to\s]+\s*(\b(?:0[1-9]|1[0-2]|[A-Za-z]{3,9})[\s/.-]*(?:20\d{2}|19\d{2})|\b(?:19|20)\d{2}|present|current|now)\b',
    re.IGNORECASE
)

# Degree extraction patterns
DEGREE_PATTERNS = [
    (r'\b(Ph\.?D\.?|Doctor of Philosophy)\b', "Ph.D."),
    (r'\b(M\.?S\.?|M\.?Tech\.?|Master of Technology|Master of Science|M\.?S\.?C\.?|M\.?C\.?A\.?|Master of Computer Applications|M\.?B\.?A\.?|Master of Business Administration)\b', "Master's"),
    (r'\b(B\.?S\.?|B\.?Tech\.?|Bachelor of Technology|Bachelor of Science|B\.?E\.?|Bachelor of Engineering|B\.?C\.?A\.?|Bachelor of Computer Applications|B\.?B\.?A\.?|Bachelor of Business Administration|Bachelor\'?s?)\b', "Bachelor's"),
    (r'\b(Diploma|Polytechnic|Associate\'?s?|A\.?S\.?)\b', "Diploma"),
    (r'\b(High School|HSC|SSC|Secondary School|12th|10th)\b', "High School")
]

# Section headers patterns (lower case keywords)
SECTION_KEYWORDS = {
    "education": [
        "education", "academic", "academics", "qualification", "qualifications",
        "academic background", "educational qualification", "studies", "schooling", "university", "college"
    ],
    "experience": [
        "experience", "employment", "work history", "professional background",
        "career", "job history", "work experience", "internship", "internships",
        "professional experience", "industry experience", "practical experience",
        "work summary", "employment history", "training & internship", "trainings & internships"
    ],
    "projects": [
        "projects", "personal project", "academic project", "key project", "portfolio",
        "academic projects", "key projects", "personal projects", "technical projects",
        "applied projects", "projects & portfolio", "project work"
    ],
    "certifications": [
        "certification", "certifications", "credential", "credentials", "certified",
        "license", "licenses", "courses", "certificates", "certifications & licenses",
        "certifications & courses", "accreditations", "certifications & training"
    ],
    "skills": [
        "skills", "technical skills", "core competencies", "technologies", "expertise",
        "technical expertise", "skills & tools", "skills & abilities", "tools & technologies",
        "programming languages", "domain skills", "key skills"
    ]
}

DEGREE_KEYWORDS_MAP = {
    "b.tech": "B.Tech",
    "btech": "B.Tech",
    "b.e.": "B.E.",
    "be": "B.E.",
    "mca": "MCA",
    "m.c.a.": "MCA",
    "bca": "BCA",
    "b.c.a.": "BCA",
    "bsc": "B.Sc",
    "b.sc": "B.Sc",
    "msc": "M.Sc",
    "m.sc": "M.Sc",
    "m.tech": "M.Tech",
    "mtech": "M.Tech",
    "phd": "Ph.D.",
    "ph.d.": "Ph.D.",
    "mba": "MBA",
    "m.b.a.": "MBA"
}

def load_taxonomy() -> Dict[str, List[str]]:
    """Loads the skills taxonomy from json config."""
    path = os.path.join(os.path.dirname(__file__), "skills_taxonomy.json")
    if os.path.exists(path):
        try:
            with open(path, "r") as f:
                return json.load(f)
        except Exception:
            pass
    return {
        "Python": ["numpy", "pandas", "scikit-learn", "django", "flask", "pytorch", "fastapi"],
        "Frontend": ["react", "vue", "angular", "javascript", "typescript", "tailwind", "html", "css"],
        "Backend": ["java", "kotlin", "spring", "node", "express", "go", "golang", "c++", "c#", ".net"],
        "Mobile": ["android", "ios", "flutter", "react native", "swift"],
        "Cybersecurity": ["owasp", "kali linux", "penetration testing", "forensics", "cryptography", "wireshark", "burp suite"],
        "DevOps": ["docker", "kubernetes", "git", "github", "ci/cd", "linux", "bash", "aws", "gcp", "azure"]
    }

def check_file_corrupted(file_path: str, extension: str) -> bool:
    """Attempts to read file structure to detect corruption."""
    try:
        ext = extension.lower()
        if ext == ".pdf":
            with pdfplumber.open(file_path) as pdf:
                if len(pdf.pages) == 0:
                    return True
            return False
        elif ext == ".docx":
            doc = docx.Document(file_path)
            _ = len(doc.paragraphs)
            return False
        elif ext == ".txt":
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                _ = f.read(1024)
            return False
    except Exception as e:
        print(f"File validation failed for {file_path}: {e}")
        return True
    return True

def extract_text(file_path: str) -> str:
    """Extracts raw text from pdf, docx, or txt files."""
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()
    text = ""
    
    if ext == ".pdf":
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    elif ext == ".docx":
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
    elif ext == ".txt":
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
            
    return sanitize_text(text)

def extract_contact_info(text: str) -> Tuple[Optional[str], Optional[str]]:
    """Extracts email and phone using regex."""
    email = None
    phone = None
    
    email_match = EMAIL_REGEX.search(text)
    if email_match:
        email = email_match.group(0)
        
    phone_match = PHONE_REGEX.search(text)
    if phone_match:
        phone = phone_match.group(0).strip()
        
    return email, phone

def extract_name_from_filename(filename: str) -> str:
    """Extracts a readable name from the resume filename."""
    name_part, _ = os.path.splitext(os.path.basename(filename))
    name_part = name_part.replace("_", " ").replace("-", " ")
    words = name_part.split()
    filtered_words = [
        w for w in words 
        if w.lower() not in {"resume", "cv", "pdf", "docx", "txt", "job", "apply", "applicant", "work", "profile"}
    ]
    return " ".join(filtered_words).title() if filtered_words else name_part.title()

def extract_name(text: str, filename: str) -> str:
    """Extracts candidate name using spaCy PERSON tags with robust heuristics and filename fallbacks."""
    filename_fallback = extract_name_from_filename(filename)

    if not nlp:
        return filename_fallback

    top_text = text[:300]
    doc = nlp(top_text)
    
    exclude_keywords = {
        "university", "college", "school", "institute", "technology", "science", "resume", "cv", "page", 
        "address", "phone", "email", "subject", "tech", "profile", "summary", "skills", "experience",
        "education", "certified", "pmp", "aws", "developer", "engineer", "designer", "architect", "management"
    }

    detected_names = []
    for ent in doc.ents:
        if ent.label_ == "PERSON":
            cleaned_ent = ent.text.strip().replace("\n", " ")
            ent_words = cleaned_ent.lower().split()
            if len(ent_words) >= 2 and len(ent_words) <= 4:
                if not any(w in exclude_keywords for w in ent_words) and not any(char.isdigit() for char in cleaned_ent):
                    detected_names.append(cleaned_ent.title())

    if detected_names:
        return detected_names[0]
        
    return filename_fallback

def extract_location(text: str) -> Optional[str]:
    """Extracts candidate location using GPE tags or city patterns."""
    if not nlp:
        return None
        
    top_text = text[:400]
    doc = nlp(top_text)
    
    locations = []
    for ent in doc.ents:
        if ent.label_ in {"GPE", "LOC"}:
            cleaned = ent.text.strip().replace("\n", " ")
            if len(cleaned) > 2 and not any(char.isdigit() for char in cleaned):
                locations.append(cleaned.title())
                
    if locations:
        return ", ".join(list(dict.fromkeys(locations))[:3])
        
    lines = [line.strip() for line in top_text.split("\n") if line.strip()]
    for line in lines:
        if "," in line and len(line) < 60:
            if "@" not in line and not any(char.isdigit() for char in line if char not in {" ", ",", "+", "-", "(", ")"}):
                return line.title()
                
    return None

def segment_text_to_sections(text: str) -> Dict[str, str]:
    """Divides resume text into chunks based on headers."""
    lines = text.split("\n")
    sections = {
        "education": [],
        "experience": [],
        "projects": [],
        "certifications": [],
        "skills": [],
        "general": []
    }
    
    current_section = "general"
    
    for line in lines:
        cleaned_line = line.strip()
        if not cleaned_line:
            continue
            
        is_header = False
        if len(cleaned_line) < 45:
            lower_line = cleaned_line.lower().replace(":", "").replace("*", "").replace("#", "").strip()
            # Strip leading bullet points if any
            lower_line = re.sub(r'^[•\-*–—\d.]+\s*', '', lower_line)
            for sect, kw_list in SECTION_KEYWORDS.items():
                if any(lower_line == kw or lower_line.startswith(kw + " ") or lower_line.endswith(" " + kw) for kw in kw_list):
                    current_section = sect
                    is_header = True
                    break
                    
        if not is_header:
            sections[current_section].append(cleaned_line)
            
    return {k: "\n".join(v) for k, v in sections.items()}

def parse_education(edu_text: str) -> List[Dict[str, Any]]:
    """Extracts degree, institution, and year from education text block."""
    entries = []
    lines = [line.strip() for line in edu_text.split("\n") if line.strip()]
    if not lines:
        return entries
        
    inst_keywords = ["university", "college", "institute", "school", "academy", "polytechnic", "vidyalaya", "patil", "iit", "nit", "bits", "nerul", "mumbai", "pune", "delhi", "engineering", "technology", "sciences"]
    
    full_edu = "\n".join(lines)
    
    # Check degrees
    detected_degree = None
    for pattern, norm in DEGREE_PATTERNS:
        if re.search(pattern, full_edu, re.IGNORECASE):
            detected_degree = norm
            break
            
    # Check graduation year
    year_match = YEAR_REGEX.search(full_edu)
    year = int(year_match.group(0)) if year_match else None
    
    # Check institution name
    institution = None
    for line in lines:
        if any(k in line.lower() for k in inst_keywords):
            clean_inst = line
            for pattern, _ in DEGREE_PATTERNS:
                clean_inst = re.sub(pattern, "", clean_inst, flags=re.IGNORECASE)
            for kw in DEGREE_KEYWORDS_MAP.keys():
                clean_inst = re.sub(r'\b' + re.escape(kw) + r'\b', "", clean_inst, flags=re.IGNORECASE)
            clean_inst = re.sub(YEAR_REGEX, "", clean_inst)
            clean_inst = clean_inst.replace("|", " ").replace("–", " ").replace("—", " ").replace("-", " ").strip()
            clean_inst = re.sub(r'^[•\-*,\s/.]+', '', clean_inst).strip()
            clean_inst = re.sub(r'\s+', " ", clean_inst).title()
            if len(clean_inst) >= 4:
                institution = clean_inst[:100]
                break
                
    if not detected_degree:
        if institution and any(k in institution.lower() for k in ["engineering", "technology", "college", "university"]):
            detected_degree = "Bachelor's"
        else:
            detected_degree = "Bachelor's" if len(lines) > 0 else None
            
    if detected_degree or institution:
        entries.append({
            "Degree": detected_degree or "Bachelor's",
            "Institution": institution or "Academic Institution",
            "Graduation_Year": year
        })
        
    return entries

def calculate_months_from_range(start_str: str, end_str: str) -> int:
    """Calculates duration in months from two date strings."""
    try:
        start_str = start_str.strip()
        end_str = end_str.strip()
        
        start_year = int(re.search(r'\b(?:19|20)\d{2}\b', start_str).group(0))
        start_month_m = re.search(r'\b(0[1-9]|1[0-2])\b', start_str)
        start_month = int(start_month_m.group(1)) if start_month_m else 1
        
        if end_str.lower() in ["present", "current", "now"]:
            end_year = 2026
            end_month = 8
        else:
            end_year_m = re.search(r'\b(?:19|20)\d{2}\b', end_str)
            end_year = int(end_year_m.group(0)) if end_year_m else 2026
            end_month_m = re.search(r'\b(0[1-9]|1[0-2])\b', end_str)
            end_month = int(end_month_m.group(1)) if end_month_m else 12
            
        months = (end_year - start_year) * 12 + (end_month - start_month) + 1
        return max(1, months)
    except Exception:
        return 12

def parse_experience(exp_text: str) -> List[Dict[str, Any]]:
    """Extracts professional and internship experience details: Company, Role, Duration, Description."""
    entries = []
    lines = [line.strip() for line in exp_text.split("\n") if line.strip()]
    if not lines:
        return entries
        
    role_keywords = ["engineer", "developer", "analyst", "manager", "intern", "internship", "specialist", "consultant", "architect", "lead", "designer", "programmer", "officer", "administrator", "researcher"]
    
    current_entry = None
    description_buffer = []
    
    for i, line in enumerate(lines):
        is_bullet = line.startswith(('•', '-', '*', '–', '—')) or re.match(r'^\d+\.', line)
        date_match = DATE_RANGE_REGEX.search(line)
        has_role = any(r in line.lower() for r in role_keywords)
        
        if (date_match or (has_role and len(line) < 80)) and not is_bullet:
            if current_entry:
                current_entry["Description"] = "\n".join(description_buffer).strip()
                entries.append(current_entry)
                description_buffer = []
                
            role = "Professional Experience"
            company = "Organization"
            duration_months = 12
            
            if date_match:
                start_d, end_d = date_match.group(1), date_match.group(2)
                duration_months = calculate_months_from_range(start_d, end_d)
                
                before_date = line[:date_match.start()].strip()
                after_date = line[date_match.end():].strip()
                
                # Check which side has role keywords
                if any(r in after_date.lower() for r in role_keywords):
                    role = after_date
                    company = before_date or "Organization"
                elif any(r in before_date.lower() for r in role_keywords):
                    role = before_date
                    company = after_date or "Organization"
                else:
                    role = after_date or before_date or "Specialist / Intern"
                    company = before_date or after_date or "Organization"
            else:
                if " at " in line:
                    parts = line.split(" at ")
                    role = parts[0].strip()
                    company = parts[1].strip()
                elif " | " in line:
                    parts = line.split(" | ")
                    role = parts[0].strip()
                    company = parts[1].strip()
                elif " - " in line:
                    parts = line.split(" - ")
                    role = parts[0].strip()
                    company = parts[1].strip()
                else:
                    role = line
                    company = "Organization"
                    
            role = re.sub(r'^[•\-*,\s/.]+', '', role).strip()
            company = re.sub(r'^[•\-*,\s/.]+', '', company).strip()
            
            current_entry = {
                "Role": role.title()[:100] if role else "Engineer / Intern",
                "Company": company.title()[:100] if company else "Organization",
                "Duration_Months": duration_months,
                "Description": "",
                "Is_Relevant": False
            }
        else:
            cleaned_bullet = re.sub(r'^[•\-*–—\d.]+\s*', '', line).strip()
            if cleaned_bullet:
                if current_entry:
                    description_buffer.append(cleaned_bullet)
                elif i == 0:
                    current_entry = {
                        "Role": "Project / Industrial Experience",
                        "Company": "Applied Practice",
                        "Duration_Months": 12,
                        "Description": "",
                        "Is_Relevant": False
                    }
                    description_buffer.append(cleaned_bullet)
                    
    if current_entry:
        current_entry["Description"] = "\n".join(description_buffer).strip()
        entries.append(current_entry)
        
    return entries

def parse_skills(text: str) -> List[Dict[str, str]]:
    """Scans text for skills based on our taxonomy, mapping occurrences."""
    taxonomy = load_taxonomy()
    extracted_skills = {}
    
    text_lower = text.lower()
    
    for category, sub_skills in taxonomy.items():
        cat_esc = re.escape(category.lower())
        if re.search(r'\b' + cat_esc + r'\b', text_lower):
            evidence = ""
            for line in text.split("\n"):
                if category.lower() in line.lower():
                    evidence = line.strip()
                    break
            extracted_skills[category] = {
                "Skill": category,
                "Skill_Level": "Expert" if "expert" in text_lower or "senior" in text_lower or "lead" in text_lower else "Intermediate",
                "Evidence_Text": evidence[:250] if evidence else f"Found mention of {category}."
            }
            
        for skill in sub_skills:
            skill_esc = re.escape(skill.lower())
            if re.search(r'\b' + skill_esc + r'\b', text_lower):
                evidence = ""
                for line in text.split("\n"):
                    if skill.lower() in line.lower():
                        evidence = line.strip()
                        break
                extracted_skills[skill.title()] = {
                    "Skill": skill.title(),
                    "Skill_Level": "Intermediate",
                    "Evidence_Text": evidence[:250] if evidence else f"Found match for {skill}."
                }
                
    return list(extracted_skills.values())

def parse_projects(proj_text: str, full_text_skills: List[Dict[str, str]]) -> List[Dict[str, Any]]:
    """Extracts project details and associates them with matched technologies."""
    entries = []
    lines = [line.strip() for line in proj_text.split("\n") if line.strip()]
    if not lines:
        return entries
        
    known_techs = [s["Skill"].lower() for s in full_text_skills]
    
    curr_name = None
    curr_techs = []
    curr_bullets = []
    
    for line in lines:
        is_bullet = line.startswith(('•', '-', '*', '–', '—')) or re.match(r'^\d+\.', line)
        if not is_bullet and (len(line) < 120 or '|' in line):
            if curr_name:
                entries.append({
                    "Project_Name": curr_name[:100],
                    "Technologies": curr_techs,
                    "Description": "\n".join(curr_bullets)
                })
                curr_bullets = []
                curr_techs = []
                
            if '|' in line:
                parts = line.split('|')
                curr_name = parts[0].strip()
                tech_part = parts[1].strip()
                tech_part = re.sub(r'\b(?:0[1-9]|1[0-2])?\s*(?:20\d{2}|19\d{2})\b', '', tech_part)
                curr_techs = [t.strip().title() for t in re.split(r'[,/]+', tech_part) if t.strip() and len(t.strip()) > 1]
            else:
                curr_name = re.sub(r'^(Project:|- Project|•)\s*', '', line).strip()
        else:
            cleaned_bullet = re.sub(r'^[•\-*–—\d.]+\s*', '', line).strip()
            if cleaned_bullet:
                curr_bullets.append(cleaned_bullet)
                
    if curr_name:
        # Match any known techs in description if curr_techs is empty
        if not curr_techs and known_techs:
            combined = (curr_name + " " + " ".join(curr_bullets)).lower()
            matched = [t.title() for t in known_techs if re.search(r'\b' + re.escape(t) + r'\b', combined)]
            curr_techs = list(set(matched))
            
        entries.append({
            "Project_Name": curr_name[:100],
            "Technologies": curr_techs,
            "Description": "\n".join(curr_bullets)
        })
        
    return entries

def parse_certifications(cert_text: str) -> List[Dict[str, str]]:
    """Extracts certification names and issuing organizations cleanly from bulleted/delimited text."""
    entries = []
    if not cert_text.strip():
        return entries
        
    raw_items = re.split(r'[•\u2022\n]+', cert_text)
    
    known_orgs = [
        "Deloitte", "Tata Forage", "Tata", "HP LIF", "HP", "Learn Prompting", "Coursera", "Udemy",
        "AWS", "Google", "Microsoft", "Oracle", "Cisco", "CompTIA", "HackerRank", "Stanford",
        "DeepLearning.AI", "Forage", "PMI", "Red Hat", "Scrum Alliance"
    ]
    
    for item in raw_items:
        item = item.strip()
        if not item or len(item) < 4:
            continue
            
        item = re.sub(r'^[•\-*–—\d.]+\s*', '', item).strip()
        
        # Check for delimiters: en-dash, em-dash, hyphen, pipe, by, from
        parts = re.split(r'\s+[–—\-|]\s+|\s+by\s+|\s+from\s+|\s+at\s+', item)
        if len(parts) == 2:
            name = parts[0].strip()
            org = parts[1].strip()
        else:
            name = item
            org = "Verified Credential"
            for ko in known_orgs:
                if ko.lower() in item.lower():
                    org = ko
                    break
                    
        entries.append({
            "Certification_Name": name[:100],
            "Issuing_Org": org[:60]
        })
        
    return entries

def parse_resume_full(file_path: str, filename: str) -> Dict[str, Any]:
    """Orchestrates full file loading, metadata extraction, chunk parsing and normalization."""
    raw_text = extract_text(file_path)
    if not raw_text.strip():
        raise ValueError("Document contains no readable text content.")
        
    email, phone = extract_contact_info(raw_text)
    name = extract_name(raw_text, filename)
    location = extract_location(raw_text)
    
    sections = segment_text_to_sections(raw_text)
    
    skills = parse_skills(raw_text)
    education = parse_education(sections["education"])
    experience = parse_experience(sections["experience"])
    projects = parse_projects(sections["projects"], skills)
    certifications = parse_certifications(sections["certifications"])
    
    return {
        "Name": name,
        "Email": email,
        "Phone": phone,
        "Location": location,
        "skills": skills,
        "experiences": experience,
        "educations": education,
        "projects": projects,
        "certifications": certifications,
        "Raw_Text": raw_text
    }
